from django.shortcuts import render
from rest_framework import viewsets
from assess.serializers import StudyYearSerializer, ClassGroupSerializer, ClassGroupExtraSerializer,  StudyPeriodSerializer, SummativeWorkSerializer, \
    WorkAssessmentSerializer, WorkCriteriaMarkSerializer, ProfileStudentSerializer, WorkGroupDateItemSerializer, \
    PeriodAssessmentSerializer, StudentWorkSerializer, ReportPeriodSerializer, StudentReportTeacherSerializer, ReportTeacherSerializer, \
    EventTypeSerializer, EventParticipationSerializer, StudentReportMentorSerializer, ReportMentorSerializer, ClassGroupStudentsSerializer, \
    WorkLoadSerializer, TextTranslateSerailizer, GenerateReportSerailizer, StudentReportPsychologistSerializer, ReportPsychologistSerializer, \
        ProfileTeacherSerializer, SubjectSerializer, WorkLoadSubjectSerializer, ClassGroupForReportTeacherSerializer, ClassGroupForReportPsychoSerializer, \
        ClassYearSerializer, ClassGroupForReportMentorSerializer, ClassGroupForMentorSerializer
 
from assess.models import StudyYear, ClassGroup, StudyPeriod, SummativeWork, WorkAssessment, WorkCriteriaMark, WorkGroupDate, PeriodAssessment, \
    ReportPeriod, ReportTeacher, EventType, EventParticipation, ReportMentor, WorkLoad, ReportPsychologist, ReportAchievements
from member.models import ProfileTeacher, ProfileStudent, User, Department
from curriculum.models import ClassYear, Subject, Criterion
from django.http.response import StreamingHttpResponse

from rest_framework.response import Response
from rest_framework.views import APIView
from django.db.models import Q

import requests
from datetime import datetime

from docxtpl import DocxTemplate
from htmldocx import HtmlToDocx
from docx import Document
from docx.shared import Inches, Pt

import io
from django.conf import settings
import os
import json

from googletrans import Translator
import openai


class StudyYearViewSet(viewsets.ModelViewSet):
    queryset = StudyYear.objects.all()
    serializer_class = StudyYearSerializer


class ClassGroupViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all()
    default_serializer_class = ClassGroupSerializer
    serializer_classes = {
        'retrieve': ClassGroupExtraSerializer,
        'update': ClassGroupExtraSerializer,
    }
    def get_queryset(self):
        class_year = self.request.query_params.get("class_year", None)
        study_year = self.request.query_params.get("study_year", None)
        level = self.request.query_params.get("level", None)
        program = self.request.query_params.get("program", None)
        teacher = self.request.query_params.get("teacher", None)
        groups = ClassGroup.objects.all()
        if class_year:
            groups = groups.filter(class_year=class_year)
        if study_year:
            groups = groups.filter(study_year=study_year)
        if level:
            groups = groups.filter(class_year__level=level)
        if program:
            groups = groups.filter(class_year__program=program)
        if teacher:
            groups = groups.filter(workload__teacher=teacher)
        return groups
    def get_serializer_class(self):
        return self.serializer_classes.get(self.action, self.default_serializer_class)


class ClassGroupStudentsViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all()
    serializer_class = ClassGroupStudentsSerializer
    def get_queryset(self):
        class_year = self.request.query_params.get("class_year", None)
        study_year = self.request.query_params.get("study_year", None)
        program = self.request.query_params.get("program", None)
        teacher = self.request.query_params.get("teacher", None)
        groups = ClassGroup.objects.all()
        if class_year:
            groups = groups.filter(class_year=class_year)
        if study_year:
            groups = groups.filter(study_year=study_year)
        if program:
            groups = groups.filter(class_year__program=program)
        if teacher:
            groups = groups.filter(workload__teacher=teacher)
        return groups

# Набор CRUD-методов для работы с моделью Студенты
class StudentViewSet(viewsets.ModelViewSet):
    queryset = ProfileStudent.objects.all()
    serializer_class = ProfileStudentSerializer
    def get_queryset(self):
        group = self.request.query_params.get("class", None)
        year = self.request.query_params.get("year", None)
        students = ProfileStudent.objects.all()
        if group:
            print(f"Get-запрос class: {group}")
            students = students.filter(groups__in=[group])
        if year:
            print(f"Get-запрос year: {year}")
            students = students.filter(groups__class_year__in=[year])
        print(f"Ответ от сервера: {students}")
        return students

class StudyPeriodViewSet(viewsets.ModelViewSet):
    queryset = StudyPeriod.objects.all()
    serializer_class = StudyPeriodSerializer
    def get_queryset(self):
        study_period = StudyPeriod.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        program = self.request.query_params.get("program", None)
        if study_year:
            study_period = study_period.filter(study_year=study_year)
        if program:
            study_period = study_period.filter(class_year__program__in=[program]).distinct()
        return study_period


class SummativeWorkViewSet(viewsets.ModelViewSet):
    queryset = SummativeWork.objects.all()
    serializer_class = SummativeWorkSerializer
    def get_queryset(self):
        summative_work = SummativeWork.objects.all()
        period = self.request.query_params.get("period", None)
        teacher = self.request.query_params.get("teacher", None)
        group = self.request.query_params.get("group", None)
        subject = self.request.query_params.get("subject", None)
        if period:
            summative_work = summative_work.filter(period=period)
        if teacher:
            summative_work = summative_work.filter(teacher=teacher)
        if group:
            summative_work = summative_work.filter(groups__in=[group]).distinct()
        if subject:
            summative_work = summative_work.filter(subject=subject)
        print(f"[SummativeWork] Ответ от сервера: {summative_work}")
        return summative_work

class WorkGroupDateItemViewSet(viewsets.ModelViewSet):
    queryset = WorkGroupDate.objects.all()
    serializer_class = WorkGroupDateItemSerializer
    def get_queryset(self):
        workgroup = WorkGroupDate.objects.all()
        work = self.request.query_params.get("work", None)
        if work:
            workgroup = workgroup.filter(student__group=work)
        return workgroup

class WorkAssessmentViewSet(viewsets.ModelViewSet):
    queryset = WorkAssessment.objects.all()
    serializer_class = WorkAssessmentSerializer
    def get_queryset(self):
        work_assessment = WorkAssessment.objects.all()
        group = self.request.query_params.get("class", None)
        if group:
            # print(f"[WorkAssessment] Get-запрос period: {group}")
            work_assessment = work_assessment.filter(student__group=group)
        # print(f"[WorkAssessment] Ответ от сервера: {work_assessment}")
        return work_assessment
    def update(self, request, pk=None, *args, **kwargs):
        print('Переданные данные: ', request.data)
        return super().update(request, pk=None, *args, **kwargs)

class WorkAssessmentGiveMarksAPIView(APIView):
    def post(self, request):
        if 'marks' in request.data:
            marks = request.data['marks']
            print(marks)
            queryset = WorkAssessment.objects.filter(id__in=marks.keys())
            for element in queryset:
                if not element.grade and type(marks[str(element.id)]) == int:
                    element.grade = marks[str(element.id)]
                    element.save()
            print(queryset)
        return Response({'result': WorkAssessmentSerializer(queryset, many=True).data})


class WorkCriteriaMarkViewSet(viewsets.ModelViewSet):
    queryset = WorkCriteriaMark.objects.all()
    serializer_class = WorkCriteriaMarkSerializer

class PeriodAssessmentViewSet(viewsets.ModelViewSet):
    queryset = PeriodAssessment.objects.all()
    serializer_class = PeriodAssessmentSerializer
    def get_queryset(self):
        period_assessment = PeriodAssessment.objects.all()
        student = self.request.query_params.get("student", None)
        period = self.request.query_params.get("period", None)
        subject = self.request.query_params.get("subject", None)
        year = self.request.query_params.get("year", None)
        if student:
            period_assessment = period_assessment.filter(student=student)
        if period:
            period_assessment = period_assessment.filter(period=period)
        if subject:
            period_assessment = period_assessment.filter(subject=subject)
        if year:
            period_assessment = period_assessment.filter(year=year)
        return period_assessment
    
class StudentWorkViewSet(viewsets.ModelViewSet):
    queryset = ProfileStudent.objects.all()
    serializer_class = StudentWorkSerializer
    def get_queryset(self):
        group = self.request.query_params.get("group", None)
        students = ProfileStudent.objects.all()
        if group:
            students = students.filter(groups__in=[group]).distinct()
        return students
    def get_serializer_context(self, *args, **kwargs):
        period = self.request.query_params.get("period", None)
        group = self.request.query_params.get("group", None)
        class_year = self.request.query_params.get("class_year", None)
        subject = self.request.query_params.get("subject", None)
        context = super().get_serializer_context()
        if period:
            context.update({"period": period})
        if group:
            context.update({"group": group})
        if group:
            context.update({"class_year": class_year})
        if subject:
            context.update({"subject": subject})
        return context

def get_marks_for_period_dnevnik(token, group, period, date):
    url_dnevnik_api = 'https://api.dnevnik.ru/v2/'
    headers = {
        'Access-Token': f'{token}',
        'Content-Type': 'application/json'
        }
    url = f'{url_dnevnik_api}edu-groups/{group}/reporting-periods/{period}/avg-marks/{date}'
    return requests.get(url, headers=headers)

class AssessmentJournalAPIView(APIView):
    def get(self, request):
        period_id = request.query_params.get("period", None)
        group_id = request.query_params.get("group", None)
        subject_id = request.query_params.get("subject", None)
        period = StudyPeriod.objects.filter(id=period_id).first()
        subject = Subject.objects.filter(id=subject_id).first()
        group = ClassGroup.objects.filter(id=group_id).first()
        return Response({
            'period': StudyPeriodSerializer(period).data,
            'group': ClassGroupSerializer(group).data,
            'subject': SubjectSerializer(subject).data,
            })

class AssessmentDnevnikAPIView(APIView):
    def get(self, request):
        period_dnevnik = request.query_params.get("period_dnevnik", None)
        group_dnevnik = request.query_params.get("group_dnevnik", None)
        subject_dnevnik = request.query_params.get("subject_dnevnik", None)
        user = request.query_params.get("user", None)
        user_queryset = User.objects.filter(id=user).first()
        response = get_marks_for_period_dnevnik(user_queryset.access_token_dnevnik,
                                                group_dnevnik,
                                                period_dnevnik,
                                                datetime.now().strftime("%Y-%m-%d"))
        student_mark = {}
        result = False
        if response.status_code == 200:
            for data in response.json():
                marks = list(filter(lambda mark: mark['subject_str'] == subject_dnevnik, data['per-subject-averages']))
                if marks:
                    student_mark[data['person_str']] = list(map(lambda mark: float(mark['avg-mark-value'].replace(',','.')), marks))[0]
            result = True
        elif response.json()['type'] == 'invalidToken':
            print('Ошибка токена доступа. Токен сброшен')
            user_queryset.access_token_dnevnik = ''
            user_queryset.save()
        else:
            print(f'Ошибка запроса: {response.json()}')
        return Response({
            'result': result,
            'marks': student_mark,
            })

# Перевод текста на выбранный язык
class TextTranslateAPIView(APIView):
    serializer_class = TextTranslateSerailizer
    def post(self, request, *args, **kwargs):
        serializer_class = TextTranslateSerailizer(data=request.data)
        if serializer_class.is_valid():
            data = serializer_class.validated_data
            text = data.get('text')
            language = data.get('language')
            translator = Translator()
            result = translator.translate(text, dest=language)
            return Response({"result": "success", "text": result.text})
        return Response({"result": "failed"})
    
# Запрос на генерацию текста
class GenerateReportAPIView(APIView):
    serializer_class = GenerateReportSerailizer
    def post(self, request, *args, **kwargs):
        serializer_class = GenerateReportSerailizer(data=request.data)
        if serializer_class.is_valid():
            data = serializer_class.validated_data
            print(data)
            model_engine = "text-davinci-003"
            openai.api_key = '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'
            if data['final_grade'] == 5:
                mood = 'хороших'
            elif data['final_grade'] == 4:
                mood = 'средних'
            else:
                mood = 'плохих'
            request_text = f"Напиши на русском языке отзыв по студенту по имени {data['name']} по {data['subject']}, который достиг {mood} результатов: {data['text']}"
            completion = openai.Completion.create(
                engine=model_engine,
                prompt=request_text,
                max_tokens=1024,
                temperature=0.7,
                top_p=1,
                frequency_penalty=0,
                presence_penalty=0
            )
            result = completion.choices[0].text
            if result.startswith('.'):
                result = result[1:]
            print(completion.choices[0].text)
            # print(request_text)
            return Response({"result": "success", "text": result.strip()})
        return Response({"result": "failed"})

class ReportTeacherJournalAPIView(APIView):
    def get(self, request):
        group_id = request.query_params.get("group", None)
        period_id = request.query_params.get("period", None)
        subject_id = request.query_params.get("subject", None)
        author_id = request.query_params.get("author", None)
        data = {}
        if period_id:
            period = ReportPeriod.objects.filter(id=period_id).first()
            data['period'] = ReportPeriodSerializer(period).data
        if group_id:
            group = ClassGroup.objects.filter(id=group_id).first()
            data['group'] = ClassGroupStudentsSerializer(group).data
        if subject_id:
            subject = Subject.objects.filter(id=subject_id).first()
            data['subject'] = SubjectSerializer(subject).data
        if author_id:
            author = ProfileTeacher.objects.filter(id=author_id).first()
            data['author'] = ProfileTeacherSerializer(author).data
        types = EventType.objects.all()
        data['event_types'] = EventTypeSerializer(types, many=True).data
        return Response(data)

class ReportPsychologistJournalAPIView(APIView):
    def get(self, request):
        group_id = request.query_params.get("group", None)
        period_id = request.query_params.get("period", None)
        period = ReportPeriod.objects.filter(id=period_id).first()
        group = ClassGroup.objects.filter(id=group_id).first()
        types = EventType.objects.all()
        return Response({
            'group': ClassGroupSerializer(group).data,
            'period': ReportPeriodSerializer(period).data,
            'event_types': EventTypeSerializer(types, many=True).data
            })

class ReportMentorJournalAPIView(APIView):
    def get(self, request):
        group_id = request.query_params.get("group", None)
        period_id = request.query_params.get("period", None)
        period = ReportPeriod.objects.filter(id=period_id).first()
        group = ClassGroup.objects.filter(id=group_id).first()
        types = EventType.objects.all()
        subjects = Subject.objects.filter(subject_year__years__group=group).distinct()
        return Response({
            'group': ClassGroupForReportMentorSerializer(group).data,
            'period': ReportPeriodSerializer(period).data,
            'subjects': SubjectSerializer(subjects, many=True).data,
            'event_types': EventTypeSerializer(types, many=True).data
            })

class ReportPeriodViewSet(viewsets.ModelViewSet):
    queryset = ReportPeriod.objects.all()
    serializer_class = ReportPeriodSerializer
    def get_queryset(self):
        report_period = ReportPeriod.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        if study_year:
            report_period = report_period.filter(study_year=study_year)
        return report_period
    
# Создание, редактирование, удаление репортов учителя  
class ReportTeacherViewSet(viewsets.ModelViewSet):
    queryset = ReportTeacher.objects.all()
    serializer_class = ReportTeacherSerializer
    def get_queryset(self):
        group_id = self.request.query_params.get("group", None)
        period_id = self.request.query_params.get("period", None)
        subject_id = self.request.query_params.get("subject", None)
        author_id = self.request.query_params.get("author", None)
        report_teacher = ReportTeacher.objects.all()
        if group_id:
            report_teacher = report_teacher.filter(student__groups__in=[group_id]).distinct()
        if period_id:
            report_teacher = report_teacher.filter(period=period_id)
        if subject_id:
            report_teacher = report_teacher.filter(subject=subject_id)
        if author_id:
            report_teacher = report_teacher.filter(author=author_id)
        return report_teacher
    def create(self, request, *args, **kwargs): 
        # Если пришли данные с названием bulk_create, то срабатывает алгоритм массового создания записей репортов
        data = request.data.pop('bulk_create', None)
        if data:
            reports = ReportTeacher.objects.filter(student__groups__id=data['group_id'], subject=data['subject_id'], period=data['period_id'], year=data['year_id'], author=data['author_id'])
            print(f"Отфильтровано записей: {reports}")
            reports_delete = reports.filter(~Q(id__in=[item['id'] for item in data['students'] if 'id' in item]))
            print(f"Записи на удаление: {reports_delete}")
            achievements_delete = ReportAchievements.objects.filter(teacher_reports__in=reports_delete)
            print(f"Достижения на удаление: {achievements_delete}")
            achievements_delete.delete()
            reports_delete.delete()
            subject = Subject.objects.filter(pk=data['subject_id']).first()
            period = ReportPeriod.objects.filter(pk=data['period_id']).first()
            year = ClassYear.objects.filter(pk=data['year_id']).first()
            new_students = []
            for item in data['students']:
                if item['id'] is None:
                    new_students.append(ReportTeacher(
                        student=ProfileStudent.objects.filter(pk=item['student_id']).first(),
                        subject=subject,
                        period=period,
                        year=year,
                        author = self.request.user.teacher))
            print(f"Записи на создание: {new_students}")
            created_students = ReportTeacher.objects.bulk_create(new_students)
            serializer = self.get_serializer(instance=created_students, many=True)
            return Response({'result': 'success', 'reports': serializer.data})
        return super().create(request, *args, **kwargs)
    def get_serializer_context(self, *args, **kwargs):
        period = self.request.query_params.get("period", None)
        class_year = self.request.query_params.get("class_year", None)
        subject = self.request.query_params.get("subject", None)
        author = self.request.query_params.get("author", None)
        context = super().get_serializer_context()
        if period:
            context.update({"period": period})
        if subject:
            context.update({"subject": subject})
        if author:
            context.update({"author": author})
        if class_year:
            context.update({"class_year": class_year})
        return context

# Получение списка студентов со сводной информацией по итоговым оценкам по предмету и репортам учителя
class StudentReportTeacherViewSet(viewsets.ModelViewSet):
    queryset = ProfileStudent.objects.all()
    serializer_class = StudentReportTeacherSerializer
    def get_queryset(self):
        group = self.request.query_params.get("group", None)
        class_year = self.request.query_params.get("class_year", None)
        students = ProfileStudent.objects.all()
        if group:
            students = students.filter(groups__in=[group]).distinct()
        if class_year:
            students = students.filter(groups__class_year__in=[class_year]).distinct()
        return students
    def get_serializer_context(self, *args, **kwargs):
        period = self.request.query_params.get("period", None)
        class_year = self.request.query_params.get("class_year", None)
        subject = self.request.query_params.get("subject", None)
        author = self.request.query_params.get("author", None)
        context = super().get_serializer_context()
        if period:
            context.update({"period": period})
        if subject:
            context.update({"subject": subject})
        if author:
            context.update({"author": author})
        if class_year:
            context.update({"class_year": class_year})
        return context

class StudentReportPsychologistViewSet(viewsets.ModelViewSet):
    queryset = ProfileStudent.objects.all()
    serializer_class = StudentReportPsychologistSerializer
    def get_queryset(self):
        group = self.request.query_params.get("group", None)
        students = ProfileStudent.objects.all()
        if group:
            students = students.filter(groups__in=[group]).distinct()
        return students

class ReportPsychologistViewSet(viewsets.ModelViewSet):
    queryset = ReportPsychologist.objects.all()
    serializer_class = ReportPsychologistSerializer

class StudentReportMentorViewSet(viewsets.ModelViewSet):
    queryset = ProfileStudent.objects.all()
    serializer_class = StudentReportMentorSerializer
    def get_queryset(self):
        group = self.request.query_params.get("group", None)
        class_year = self.request.query_params.get("class_year", None)
        students = ProfileStudent.objects.all()
        if group:
            students = students.filter(groups__in=[group]).distinct()
        if class_year:
            students = students.filter(groups__class_year__in=[class_year]).distinct()
        return students
    def get_serializer_context(self, *args, **kwargs):
        period = self.request.query_params.get("period", None)
        class_year = self.request.query_params.get("class_year", None)
        study_year = self.request.query_params.get("study_year", None)
        author = self.request.query_params.get("author", None)
        context = super().get_serializer_context()
        if period:
            context.update({"period": period})
        if study_year:
            context.update({"study_year": study_year})
        if author:
            context.update({"author": author})
        if class_year:
            context.update({"class_year": class_year})
        return context

class ReportMentorViewSet(viewsets.ModelViewSet):
    queryset = ReportMentor.objects.all()
    serializer_class = ReportMentorSerializer

class EventTypeViewSet(viewsets.ModelViewSet):
    queryset = EventType.objects.all()
    serializer_class = EventTypeSerializer

class EventParticipationViewSet(viewsets.ModelViewSet):
    queryset = EventParticipation.objects.all()
    serializer_class = EventParticipationSerializer
    def get_queryset(self):
        event_participation = EventParticipation.objects.all()
        type = self.request.query_params.get("type", None)
        level = self.request.query_params.get("level", None)
        student = self.request.query_params.get("student", None)
        if type:
            event_participation = event_participation.filter(type=type)
        if level:
            event_participation = event_participation.filter(level=level)
        if student:
            event_participation = event_participation.filter(teacher_reports__student__in=[student], mentor_reports__student__in=[student])
        return event_participation
    
class WorkLoadViewSet(viewsets.ModelViewSet):
    queryset = WorkLoad.objects.all()
    serializer_class = WorkLoadSerializer
    def get_queryset(self):
        workload = WorkLoad.objects.all()
        teacher = self.request.query_params.get("teacher", None)
        study_year = self.request.query_params.get("study_year", None)
        if teacher:
            workload = workload.filter(teacher=teacher)
        if study_year:
            workload = workload.filter(study_year=study_year)
        return workload

def get_criterion_rus(criterion_count, criterion_summ):
    GRADES = { 1: [3, 5, 7], 2: [6, 10, 14], 3: [8, 14, 20], 4: [11, 19, 28] }
    if criterion_count not in GRADES:
        return "N/A"
    if criterion_summ >= GRADES[criterion_count][2]:
        return 5
    elif criterion_summ < GRADES[criterion_count][2] and criterion_summ >= GRADES[criterion_count][1]:
        return 4
    elif criterion_summ < GRADES[criterion_count][1] and criterion_summ >= GRADES[criterion_count][0]:
        return 3
    elif criterion_summ < GRADES[criterion_count][0] and criterion_summ > 0:
        return 2
    else:
        return '-'

# Экспорт в docx-файл
class ExportReportMentorDOCXAPIView(APIView):
    def get(self, request, *args, **kwargs):
        group_id = self.request.query_params.get("group", None)
        period_id = self.request.query_params.get("period", None)
        student_id = self.request.query_params.get("student", None)
        report_id = self.request.query_params.get("report", None)
        student = ProfileStudent.objects.filter(pk=student_id).first()
        teacher_reports = ReportTeacher.objects.filter(student=student,
                                                        period__id=period_id,
                                                        year__group=group_id)
        psycho_report = ReportPsychologist.objects.filter(student=student,
                                                            period__id=period_id,
                                                            year__group=group_id).first()
        mentor_report = ReportMentor.objects.filter(pk=report_id).first()
        # create an empty document object
        document = self.build_document(mentor_report, teacher_reports, psycho_report)
        # save document info
        buffer = io.BytesIO()
        document.save(buffer)  # save your memory stream
        buffer.seek(0)  # rewind the stream
        response = StreamingHttpResponse(
            streaming_content=buffer,  # use the stream's content
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment;filename={student.user.username}.docx'
        response["Content-Encoding"] = 'UTF-8'
        return response
    def build_document(self, mentor_report, teacher_reports, psycho_report):
        # document = Document(os.path.join(settings.BASE_DIR, 'reports', 'temp_report_myp.docx'))
        document = Document()
        section = document.sections[0]
        section.header_distance = Pt(10)
        section.footer_distance = Pt(10)
        # print(document._body._body.xml)
        document.styles['Normal'].paragraph_format.space_after = Pt(0)
        parser = HtmlToDocx()
        if mentor_report:
            heading_document = document.add_heading(f"Academic year report {mentor_report.period.study_year.name}", level=1)
            heading_document.alignment = 1
            document.add_paragraph(f"\nStudent name: {mentor_report.student.user.last_name} {mentor_report.student.user.first_name}").paragraph_format.space_after = None
            group = mentor_report.year.group.filter(study_year=mentor_report.period.study_year, students=mentor_report.student).first()
            document.add_paragraph(f"Grade: {group.group_name} класс ({mentor_report.year.year_ib})")
            document.add_paragraph(f"Advisor: {mentor_report.author.user.last_name} {mentor_report.author.user.first_name} {mentor_report.author.user.middle_name}")
            document.add_heading(f"REFLECTIONS", level=2)
            # Блок репорта наставника
            title_mentor = document.add_heading(f"Homeroom Advisor (Наставник)", level=3)
            document.add_paragraph()
            for element in parser.parse_html_string(mentor_report.text)._element:
                title_mentor._p.addnext(element)
            # Блок участия в мероприятиях
            title_mentor = document.add_heading(f"Участие в мероприятиях", level=3)
            events_queryset = EventParticipation.objects.filter(Q(teacher_reports__student=mentor_report.student, 
                                                            teacher_reports__period=mentor_report.period,
                                                            teacher_reports__year=mentor_report.year) | 
                                                            Q(psycho_reports__student=mentor_report.student,
                                                              psycho_reports__period=mentor_report.period,
                                                              psycho_reports__year=mentor_report.year) | 
                                                            Q(mentor_reports__student=mentor_report.student,
                                                              mentor_reports__period=mentor_report.period,
                                                              mentor_reports__year=mentor_report.year))
            for event in events_queryset:
                document.add_paragraph(f"{event.title} ({event.get_level_display()}, {event.type.name}): {event.result}")
            # Блок репорта психолога
            title_psychologist = document.add_heading(f"Psychologist (Психолог)", level=3)
            if psycho_report:
                for element in parser.parse_html_string(psycho_report.text)._element:
                    title_psychologist._p.addnext(element)
            # Блок репортов учителей-предметников
            document.add_heading(f"Teachers (Учителя)", level=3)
            document.add_paragraph()
            departments = Department.objects.all()
            for department in departments:
                reports = teacher_reports.filter(subject__department=department)
                if reports:
                    if group.program == 'myp':
                        table = document.add_table(rows=1, cols=6, style="Table Grid")
                    elif group.program == 'dp':
                        table = document.add_table(rows=1, cols=3, style="Table Grid")
                    else:
                        table = document.add_table(rows=1, cols=2, style="Table Grid")
                    table.rows[0].cells[0].text = f"{department.name.upper()}"
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
                    table.rows[0].cells[0].paragraphs[0].runs[0].font.size = Pt(16)
                    count = 1
                    for j, report in enumerate(reports):
                        table.add_row()
                        table.rows[j + count].cells[0].text = f"Предмет"
                        if group.program == 'myp':
                            table.rows[j + count].cells[1].text = f"Критерии оценивания"
                            table.rows[j + count].cells[1].width = Inches(200.0)
                            table.rows[j + count].cells[2].text = f"Баллы"
                            table.rows[j + count].cells[3].text = f"Сумма баллов"
                            table.rows[j + count].cells[4].text = f"Оценка"
                            table.rows[j + count].cells[5].text = f"Итог"
                        elif group.program == 'dp':
                            table.rows[j + count].cells[1].text = f"Final Grade\nNon-IB term grade"
                            table.rows[j + count].cells[2].text = f"Final Grade\nIB term grade"
                        else:
                            table.rows[j + count].cells[1].text = f"Final Grade"
                        table.add_row()
                        count += 1
                        table.rows[j + count].cells[0].text = f"{report.subject.name_rus}"
                        if group.program == 'myp':
                            criteria = report.report_criteria.all()
                            if criteria:
                                criterion_summ = 0
                                for i, item in enumerate(criteria):
                                    table.rows[j + i + count].cells[1].text = f"{item.criterion.letter}. {item.criterion.name_eng}"
                                    table.rows[j + i + count].cells[2].text = f"{item.mark}"
                                    table.add_row()
                                    criterion_summ += item.mark
                                table.rows[j + count].cells[3].text = f"{criterion_summ}/{len(criteria) * 8}"
                                table.rows[j + count].cells[4].text = f"{get_criterion_rus(len(criteria), criterion_summ)}"
                                table.rows[j + count].cells[5].text = f"{report.final_grade}"
                                table.cell(j + count, 0).merge(table.cell(j + count + len(criteria) - 1, 0))
                                table.cell(j + count, 3).merge(table.cell(j + count + len(criteria) - 1, 3))
                                table.cell(j + count, 4).merge(table.cell(j + count + len(criteria) - 1, 4))
                                table.cell(j + count, 5).merge(table.cell(j + count + len(criteria) - 1, 5))
                            else:
                                table.add_row()
                                count += 1
                            count += len(criteria)
                        elif group.program == 'dp':
                            table.rows[j + count].cells[1].text = f"{report.final_grade}"
                            table.rows[j + count].cells[2].text = f"{report.final_grade_ib}"
                            table.add_row()
                            count += 1
                        else:
                            table.rows[j + count].cells[1].text = f"{report.final_grade}"
                            table.add_row()
                            count += 1
                        report_cell = table.cell(j + count, 0).paragraphs[0]
                        report_cell.add_run(f'Комментарии учителя ({report.author.user.last_name} {report.author.user.first_name} {report.author.user.middle_name}):').bold = True
                        if report.text:
                            for element in parser.parse_html_string(report.text)._element:
                                report_cell._p.addnext(element)
                        # self.delete_paragraph(report_cell)
                        if group.program == 'myp':
                            table.cell(0, 0).merge(table.cell(0, 5))
                            table.cell(j + count, 0).merge(table.cell(j + count, 5))
                        elif group.program == 'dp':
                            table.cell(0, 0).merge(table.cell(0, 2))
                            table.cell(j + count, 0).merge(table.cell(j + count, 2))
                        else:
                            table.cell(0, 0).merge(table.cell(0, 1))
                            table.cell(j + count, 0).merge(table.cell(j + count, 1))
                    document.add_paragraph()   
        # Добавление в колонтитул изображения
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.paragraph_format.left_indent = Pt(-30)
        logo_run = paragraph.add_run()
        logo_run.add_picture(os.path.join(settings.BASE_DIR, 'reports', 'header_report.png'))  
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.paragraph_format.left_indent = Pt(-80)
        logo_run = paragraph.add_run()
        logo_run.add_picture(os.path.join(settings.BASE_DIR, 'reports', 'footer_report.png'))
            
        return document
    def delete_paragraph(self, paragraph):
        temp_ = paragraph._element
        temp_.getparent().remove(temp_)
        temp_._p = temp_._element = None


def get_final_marks_dnevnik(token, student, group):
    url_dnevnik_api = 'https://api.dnevnik.ru/v2/'
    headers = {
        'Access-Token': f'{token}',
        'Content-Type': 'application/json'
        }
    url = f'{url_dnevnik_api}persons/{student}/edu-groups/{group}/final-marks'
    return requests.get(url, headers=headers)

def get_list_works_dnevnik(token, data):
    url_dnevnik_api = 'https://api.dnevnik.ru/v2/'
    headers = {
        'Access-Token': f'{token}',
        'Content-Type': 'application/json'
        }
    url = f'{url_dnevnik_api}works/many'
    print(data)
    return requests.post(url, data=json.dumps(data), headers=headers)
    
class FinalMarksDnevnikAPIView(APIView):
    def get(self, request):
        period_dnevnik = request.query_params.get("period_dnevnik", None)
        group_dnevnik = request.query_params.get("group_dnevnik", None)
        subject_dnevnik = request.query_params.get("subject_dnevnik", None)
        student_dnevnik = request.query_params.get("student_dnevnik", None)
        user = request.query_params.get("user", None)
        user_queryset = User.objects.filter(id=user).first()
        response_marks = get_final_marks_dnevnik(user_queryset.access_token_dnevnik,
                                                student_dnevnik,
                                                group_dnevnik)
        data = {
            'isResult': False,
            'isTokenValid': True,
        }
        marks = {}
        works = []
        if response_marks.status_code == 200:
            for mark in response_marks.json():
                if mark['subject_str'] == subject_dnevnik:
                    marks[mark['final-mark']['work_str']] = mark['final-mark']
                    works.append(mark['final-mark']['work_str'])
            data['isResult'] = True
            data['marks'] = marks
            response_works = get_list_works_dnevnik(user_queryset.access_token_dnevnik, works)
            if response_works.status_code == 200:
                for work in response_works.json():
                    data['marks'][work['id_str']]['work_full'] = work
            else:
                print(f'Ошибка запроса: {response_works.json()}')   
        elif response_marks.json()['type'] == 'invalidToken':
            print('Ошибка токена доступа. Токен сброшен')
            user_queryset.access_token_dnevnik = ''
            user_queryset.save()
            data['isTokenValid'] = False
        else:
            print(f'Ошибка запроса: {response_marks.json()}')
        return Response(data)
    
class WorkLoadSubjectViewSet(viewsets.ModelViewSet):
    queryset = Subject.objects.all()
    serializer_class = WorkLoadSubjectSerializer
    def get_queryset(self):
        subjects = Subject.objects.all()
        department = self.request.query_params.get("department", None)
        if department:
            subjects = subjects.filter(department=department)
        return subjects
    
    
class ClassYearForReportViewSet(viewsets.ModelViewSet):
    queryset = ClassYear.objects.all()
    serializer_class = ClassYearSerializer
    def get_queryset(self):
        class_year = ClassYear.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        teacher = self.request.query_params.get("teacher", None)
        psychologist = self.request.query_params.get("psychologist", None)
        mentor = self.request.query_params.get("mentor", None)
        level = self.request.query_params.get("level", None)
        if teacher:
            class_year = class_year.filter(group__workload__study_year=study_year, group__workload__teacher=teacher).distinct()
        if psychologist:
            class_year = class_year.filter(group__study_year=study_year, group__psychologist=psychologist).distinct()
        if mentor:
            class_year = class_year.filter(group__study_year=study_year, group__mentor=mentor).distinct()
        if level:
            class_year = class_year.filter(level=level)
        return class_year

class ClassGroupForReportTeacherViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all()
    serializer_class = ClassGroupForReportTeacherSerializer
    def get_queryset(self):
        class_group = ClassGroup.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        class_year = self.request.query_params.get("class_year", None)
        teacher = self.request.query_params.get("teacher", None)
        subject = self.request.query_params.get("subject", None)
        level = self.request.query_params.get("level", None)
        if study_year:
            class_group = class_group.filter(study_year=study_year)
        if class_year:
            class_group = class_group.filter(class_year=class_year)
        if level:
            class_group = class_group.filter(class_year__level=level)
        if subject:
            class_group = class_group.filter(workload__subject=subject).distinct()
        if teacher:
            class_group = class_group.filter(workload__teacher=teacher).distinct()
        return class_group
    def get_serializer_context(self, *args, **kwargs):
        subject = self.request.query_params.get("subject", None)
        context = super().get_serializer_context()
        if subject:
            context.update({"subject": subject})
        return context

class ClassGroupForReportPsychoViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all()
    serializer_class = ClassGroupForReportPsychoSerializer
    def get_queryset(self):
        class_group = ClassGroup.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        class_year = self.request.query_params.get("class_year", None)
        psychologist = self.request.query_params.get("psychologist", None)
        if study_year:
            class_group = class_group.filter(study_year=study_year)
        if class_year:
            class_group = class_group.filter(class_year=class_year)
        if psychologist:
            class_group = class_group.filter(psychologist=psychologist)
        return class_group

class ClassGroupForReportMentorViewSet(viewsets.ModelViewSet):
    queryset = ClassGroup.objects.all()
    serializer_class = ClassGroupForReportMentorSerializer
    def get_queryset(self):
        class_group = ClassGroup.objects.all()
        study_year = self.request.query_params.get("study_year", None)
        class_year = self.request.query_params.get("class_year", None)
        mentor = self.request.query_params.get("mentor", None)
        if study_year:
            class_group = class_group.filter(study_year=study_year)
        if class_year:
            class_group = class_group.filter(class_year=class_year)
        if mentor:
            class_group = class_group.filter(mentor=mentor)
        return class_group